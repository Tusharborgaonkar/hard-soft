import json
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from fpdf import FPDF

# Constants
FONT_PATH = r"C:\Windows\Fonts\shruti.ttf"
OUTPUT_DIR = "exports"

class PDFReport(FPDF):
    def header(self):
        self.set_font('Shruti', 'B', 16)
        # Using new_x and new_y for fpdf2 v2.5+
        self.cell(0, 10, self.report_title, align='C', new_x="LMARGIN", new_y="NEXT")
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Shruti', '', 8)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')

def load_stats():
    with open("stats_data.json", "r", encoding="utf-8") as f:
        return json.load(f)

def create_pdf_report(group_name, stats, output_path):
    pdf = PDFReport()
    pdf.report_title = f"Result Report - {group_name}"
    
    # Add Gujarati font
    if os.path.exists(FONT_PATH):
        pdf.add_font('Shruti', '', FONT_PATH)
        pdf.add_font('Shruti', 'B', FONT_PATH)
    else:
        # Fallback to Arial if Shruti not found (will not render Gujarati correctly)
        pdf.set_font('Arial', '', 12)
        
    pdf.add_page()
    pdf.set_font('Shruti', '', 12)
    
    total_resp = stats.get('1', {}).get('total_responses', 330)
    pdf.cell(0, 10, f"Total Respondents: {total_resp}", new_x="LMARGIN", new_y="NEXT")
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    for q_id, q_data in stats.items():
        pdf.set_font('Shruti', 'B', 12)
        pdf.multi_cell(0, 10, f"Q{q_id}. {q_data['question']}")
        
        pdf.set_font('Shruti', 'B', 10)
        pdf.cell(100, 8, 'Option / Answer', border=1)
        pdf.cell(40, 8, 'Count', border=1)
        pdf.cell(40, 8, 'Percentage (%)', border=1, new_x="LMARGIN", new_y="NEXT")
        
        pdf.set_font('Shruti', '', 10)
        for result in q_data['results']:
            pdf.cell(100, 8, str(result['option']), border=1)
            pdf.cell(40, 8, str(result['count']), border=1)
            pdf.cell(40, 8, f"{result['percentage']}%", border=1, new_x="LMARGIN", new_y="NEXT")
        
        pdf.ln(5)
        
        # Check for page break
        if pdf.get_y() > 250:
            pdf.add_page()

    pdf.output(output_path)
    print(f"PDF report saved: {output_path}".encode('ascii', 'ignore').decode())

def create_word_report(group_name, stats, output_path):
    doc = Document()
    title = doc.add_heading(f'Result Report - {group_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Total Respondents: 330")
    doc.add_paragraph("_" * 50)
    
    for q_id, q_data in stats.items():
        doc.add_heading(f"Q{q_id}. {q_data['question']}", level=2)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Option / Answer'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage (%)'
        
        for result in q_data['results']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(result['option'])
            row_cells[1].text = str(result['count'])
            row_cells[2].text = f"{result['percentage']}%"
            
        doc.add_paragraph("")

    doc.save(output_path)
    print(f"Word report saved: {output_path}".encode('ascii', 'ignore').decode())

def create_excel_report(all_stats, output_path):
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        for group_name, stats in all_stats.items():
            rows = []
            for q_id, q_data in stats.items():
                for result in q_data['results']:
                    rows.append({
                        "Question ID": q_id,
                        "Question": q_data['question'],
                        "Option": result['option'],
                        "Count": result['count'],
                        "Percentage": result['percentage']
                    })
            
            df = pd.DataFrame(rows)
            sheet_name = group_name[:31].replace("(", "").replace(")", "")
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    print(f"Excel report saved: {output_path}")

def run():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        
    all_stats = load_stats()
    
    for group_name, stats in all_stats.items():
        if group_name == "Final Combined": continue
        
        safe_name = group_name.replace("(", "").replace(")", "").replace(" ", "_")
        
        # Word
        word_path = os.path.join(OUTPUT_DIR, f"{safe_name}_Report.docx")
        create_word_report(group_name, stats, word_path)
        
        # PDF
        pdf_path = os.path.join(OUTPUT_DIR, f"{safe_name}_Report.pdf")
        create_pdf_report(group_name, stats, pdf_path)
        
    # Final Combined
    create_word_report("Final Combined", all_stats['Final Combined'], os.path.join(OUTPUT_DIR, "Final_Combined_Report.docx"))
    create_pdf_report("Final Combined", all_stats['Final Combined'], os.path.join(OUTPUT_DIR, "Final_Combined_Report.pdf"))
    create_excel_report(all_stats, os.path.join(OUTPUT_DIR, "Final_Combined_Report.xlsx"))

if __name__ == "__main__":
    run()
